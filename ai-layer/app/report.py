# produce PDF / Word reports
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
import os
from app.storage import list_summaries
from app.config import settings
import datetime

async def generate_daily_report(output_basename=None, since_ts=None):
    # collect summaries
    docs = await list_summaries(since_ts=since_ts, limit=200)
    if not output_basename:
        ts = datetime.datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
        output_basename = f"ai-report-{ts}"
    os.makedirs(settings.REPORT_OUTPUT_DIR, exist_ok=True)
    docx_path = os.path.join(settings.REPORT_OUTPUT_DIR, output_basename + ".docx")
    pdf_path = os.path.join(settings.REPORT_OUTPUT_DIR, output_basename + ".pdf")
    # Word
    doc = Document()
    doc.add_heading("AI News Daily Report", level=1)
    for s in docs:
        doc.add_heading(s.get("title","(no title)"), level=2)
        doc.add_paragraph(f"Source: {s.get('source','')}")
        doc.add_paragraph(s.get("summary",""))
        doc.add_paragraph("Highlights:")
        for h in s.get("highlights", []):
            doc.add_paragraph(f"- {h}")
        doc.add_paragraph("\n---\n")
    doc.save(docx_path)
    # PDF (very simple)
    c = canvas.Canvas(pdf_path, pagesize=A4)
    width, height = A4
    y = height - 40
    c.setFont("Helvetica-Bold", 16)
    c.drawString(40, y, "AI News Daily Report")
    y -= 30
    c.setFont("Helvetica", 10)
    for s in docs:
        if y < 80:
            c.showPage()
            y = height - 40
        c.setFont("Helvetica-Bold", 12)
        c.drawString(40, y, s.get("title","(no title)"))
        y -= 16
        c.setFont("Helvetica", 9)
        c.drawString(45, y, f"Source: {s.get('source','')}")
        y -= 12
        for line in s.get("summary","").split("\n"):
            c.drawString(45, y, line[:120])
            y -= 12
            if y < 80:
                c.showPage()
                y = height - 40
        y -= 8
    c.save()
    return {"docx": docx_path, "pdf": pdf_path}

