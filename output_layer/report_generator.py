# generate HTML / PDF reports
# report_generator.py
from jinja2 import Environment, FileSystemLoader
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import os
from .output_config import config

env = Environment(loader=FileSystemLoader("layer5_output/templates"))
styles = getSampleStyleSheet()

class ReportGenerator:

    @staticmethod
    def generate_pdf(summary_items: list, file_path: str):
        template = env.get_template("base_report.html")
        html = template.render(items=summary_items)

        doc = SimpleDocTemplate(file_path)
        story = [Paragraph(html, styles["Normal"])]
        doc.build(story)

    @staticmethod
    def generate_docx(summary_items: list, file_path: str):
        template = env.get_template("base_report.html")
        html = template.render(items=summary_items)

        doc = Document()
        doc.add_heading("AI News Summary", level=1)
        doc.add_paragraph(html)
        doc.save(file_path)

